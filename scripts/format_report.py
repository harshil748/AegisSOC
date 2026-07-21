#!/usr/bin/env python3
"""Reformat report.docx: sections, roman/arabic page numbers, TOC, diagrams, screenshot captions."""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC = Path("/Users/harshil/Downloads/report.docx")
OUT = Path("/Users/harshil/Downloads/report.docx")
ASSETS = ROOT / "docs" / "report-assets" / "extracted"
DIAGRAMS = ROOT / "docs" / "report-assets" / "diagrams"
DIAGRAMS.mkdir(parents=True, exist_ok=True)

STUDENT_ID = "23DCE081"
FONT = "Times New Roman"

# Estimated page numbers (Acknowledgement = roman i; Chapter 1 = arabic 1)
TOC_ENTRIES = [
    ("ACKNOWLEDGEMENT", "i", False),
    ("ABSTRACT", "ii", False),
    ("LIST OF FIGURES", "iii", False),
    ("LIST OF TABLES", "iv", False),
    ("DESCRIPTION OF COMPANY / ORGANIZATION", "v", False),
    ("CHAPTER 1  PROJECT DEFINITION", "1", True),
    ("    1.1   Introduction", "1", True),
    ("    1.2   Problem Statement", "1", True),
    ("    1.3   Project Objectives", "2", True),
    ("    1.4   Scope", "2", True),
    ("CHAPTER 2  DESCRIPTION", "3", True),
    ("    2.1   System Overview", "3", True),
    ("    2.2   Architecture Layers", "3", True),
    ("    2.3   Technology Stack", "5", True),
    ("CHAPTER 3  SOFTWARE AND HARDWARE REQUIREMENTS", "6", True),
    ("    3.1   Software Requirements", "6", True),
    ("    3.2   Hardware Requirements", "6", True),
    ("CHAPTER 4  MAJOR FUNCTIONALITY", "7", True),
    ("    4.1   Microservices", "7", True),
    ("    4.2   Key Features", "8", True),
    ("    4.3   Demo Scenarios", "8", True),
    ("CHAPTER 5  FLOW CHART", "9", True),
    ("    5.1   End-to-End Pipeline Flow", "9", True),
    ("    5.2   Microservices Architecture", "10", True),
    ("    5.3   Alert-to-Response Sequence", "11", True),
    ("CHAPTER 6  SCREENSHOTS OF PROJECT OUTPUT", "12", True),
    ("    6.1   Alert Queue Dashboard", "12", True),
    ("    6.2   Cases Management", "13", True),
    ("    6.3   Case Detail View", "14", True),
    ("    6.4   Investigation Workspace", "15", True),
    ("    6.5   Response and Approvals", "16", True),
    ("    6.6   Approval Modal (Human-in-the-Loop)", "17", True),
    ("    6.7   Audit Trail", "18", True),
    ("    6.8   Metrics Dashboard", "19", True),
    ("    6.9   Replay and Demo Console", "20", True),
    ("CHAPTER 7  LIMITATIONS OF PROJECT", "21", True),
    ("CHAPTER 8  OUTCOME", "22", True),
    ("CHAPTER 9  FUTURE ENHANCEMENT", "23", True),
    ("CONCLUSION", "24", True),
    ("REFERENCES", "25", True),
]

LOF = [
    ("5.1", "AegisSOC End-to-End Data Flow", "9"),
    ("5.2", "Microservices Architecture Diagram", "10"),
    ("5.3", "Alert-to-Response Sequence Diagram", "11"),
    ("6.1", "Alert Queue Dashboard", "12"),
    ("6.2", "Cases Management View", "13"),
    ("6.3", "Case Detail View", "14"),
    ("6.4", "Investigation Workspace with Attack-Path Graph", "15"),
    ("6.5", "Response and Approvals Panel", "16"),
    ("6.6", "Human-in-the-Loop Approval Modal", "17"),
    ("6.7", "Audit Trail with Expandable Events", "18"),
    ("6.8", "Operational Metrics Dashboard", "19"),
    ("6.9", "Replay and Demo Scenario Console", "20"),
]

LOT = [
    ("3.1", "Software Requirements", "6"),
    ("3.2", "Hardware Requirements", "6"),
    ("4.1", "Microservices and Responsibilities", "7"),
    ("4.2", "Demo Scenarios and Expected Outcomes", "8"),
]

SCREENSHOTS = [
    ("6.1", "Alert Queue Dashboard",
     "The Alert Queue presents prioritized alerts ranked by ensemble risk score. Analysts can filter by severity and status, view MITRE ATT&CK technique tags, and click any row to open the Investigation Workspace.",
     "image5.png"),
    ("6.2", "Cases Management View",
     "The Cases page lists investigation cases clustered from correlated alerts, showing severity, risk percentage, status, assignee, and last-updated timestamp for workflow tracking.",
     "image7.png"),
    ("6.3", "Case Detail View",
     "Case Detail provides a summary of the selected incident including risk score, attack story narrative, linked alerts, and mapped MITRE techniques, with a direct link to the Investigation Workspace.",
     "image8.png"),
    ("6.4", "Investigation Workspace with Attack-Path Graph",
     "The Investigation Workspace combines an interactive attack-path graph (entity nodes and typed edges), LLM triage report with groundedness score, ATT&CK mapping, and containment recommendations.",
     "image6.png"),
    ("6.5", "Response and Approvals Panel",
     "The Response page displays AI-recommended actions such as host isolation and domain blocking. Disruptive actions are flagged and require explicit analyst approval before execution.",
     "image9.png"),
    ("6.6", "Human-in-the-Loop Approval Modal",
     "The approval modal enforces human-in-the-loop governance: analysts review impact summary, toggle dry-run mode, document rationale, and confirm or cancel disruptive response actions.",
     "image10.png"),
    ("6.7", "Audit Trail with Expandable Events",
     "The Audit Trail provides an append-only log of system and analyst actions including triage report generation, alert creation, and case decisions, with expandable JSON detail per event.",
     "image11.png"),
    ("6.8", "Operational Metrics Dashboard",
     "The Metrics dashboard shows Prometheus-style operational KPIs across ingestion throughput, detection precision, case triage latency, LLM groundedness, and pending response approvals.",
     "image12.png"),
    ("6.9", "Replay and Demo Scenario Console",
     "The Replay page executes three canonical demo scenarios (phishing-ransomware, benign false-positive, repeat attacker infrastructure) to validate end-to-end pipeline behavior.",
     "image13.png"),
]


def _font(run, size=12, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _para(doc, text, *, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, spacing=1.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _font(r, size, bold)
    pf = p.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"• {text}")
    _font(r)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(3)
    return p


def _caption(doc, text):
    return _para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, before=6, after=12)


def _chapter(doc, title):
    doc.add_page_break()
    p = _para(doc, title, size=16, bold=True, before=12, after=12)
    for r in p.runs:
        r.font.all_caps = True
    return p


def _section(doc, num, title):
    return _para(doc, f"{num}   {title.upper()}", size=14, bold=True, before=12, after=6)


def _subsection(doc, num, title):
    t = title[0].upper() + title[1:] if title else title
    return _para(doc, f"{num}   {t}", size=12, bold=True, before=6, after=6)


def _pic(doc, path, width=Inches(5.8)):
    if Path(path).exists():
        doc.add_picture(str(path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_field(run, instr: str):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_el)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _set_pg_num_type(section, fmt: str, start: int):
  sectPr = section._sectPr
  pgNumType = sectPr.find(qn("w:pgNumType"))
  if pgNumType is None:
    pgNumType = OxmlElement("w:pgNumType")
    sectPr.append(pgNumType)
  pgNumType.set(qn("w:fmt"), fmt)
  pgNumType.set(qn("w:start"), str(start))


def _clear_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.clear()
    for p in section.footer.paragraphs:
        p.clear()


def _setup_footer(section, roman: bool = False):
    _clear_header_footer(section)
    # Header: Student ID (left) | optional chapter label (right)
    hp = section.header.paragraphs[0]
    hp.clear()
    r1 = hp.add_run(STUDENT_ID)
    _font(r1, 10)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)
    if not roman:
        r2 = hp.add_run("\tAegisSOC Report")
        _font(r2, 10)

    # Footer: DEPSTAR (left) | PAGE field (center) | Department (right)
    fp = section.footer.paragraphs[0]
    fp.clear()
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    r_left = fp.add_run("DEPSTAR")
    _font(r_left, 10)
    fp.add_run("\t")
    r_page = fp.add_run()
    _font(r_page, 10)
    _add_field(r_page, " PAGE ")
    fp.add_run("\t")
    r_right = fp.add_run("Department of Computer Engineering")
    _font(r_right, 10)

    if roman:
        _set_pg_num_type(section, "lowerRoman", 1)
    else:
        _set_pg_num_type(section, "decimal", 1)


def _set_margins(section):
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)


def _toc_line(doc, title, page, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(title)
    _font(r1, 12, bold)
    p.add_run("\t")
    r2 = p.add_run(page)
    _font(r2, 12)
    return p


def _add_table(doc, headers, rows, caption):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            _font(r, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    _caption(doc, caption)


def _make_diagrams():
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

    # Fig 5.1 - Pipeline flow
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 4.5); ax.axis("off")
    boxes = [
        (0.2, 2.8, "Telemetry\nSources"), (1.8, 2.8, "Ingestion"), (3.4, 2.8, "Normalization"),
        (5.0, 2.8, "Enrichment"), (6.6, 2.8, "Graph\nBuilder"), (8.2, 2.8, "Neo4j"),
        (5.0, 1.2, "Detection"), (3.2, 0.2, "Case\nMgmt"), (5.0, 0.2, "LLM\nTriage"),
        (6.8, 0.2, "Response"), (8.4, 0.2, "Analyst\nUI"),
    ]
    for x, y, lbl in boxes:
        ax.add_patch(FancyBboxPatch((x, y), 1.3, 0.75, boxstyle="round,pad=0.04", fc="#eff6ff", ec="#2563eb", lw=1.2))
        ax.text(x + 0.65, y + 0.38, lbl, ha="center", va="center", fontsize=7, fontweight="bold")
    for s, e in [((1.5,3.15),(1.8,3.15)),((3.1,3.15),(3.4,3.15)),((4.7,3.15),(5.0,3.15)),((6.3,3.15),(6.6,3.15)),((7.9,3.15),(8.2,3.15)),
                 ((5.65,2.8),(5.65,1.95)),((5.65,1.2),(4.5,0.95)),((4.5,0.58),(3.2,0.58)),((4.5,0.58),(5.0,0.58)),((5.65,0.58),(6.8,0.58)),((7.5,0.58),(8.4,0.58))]:
        ax.add_patch(FancyArrowPatch(s, e, arrowstyle="->", color="#334155", lw=1.1))
    ax.set_title("Fig. 5.1 — AegisSOC End-to-End Data Flow", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(DIAGRAMS / "fig_5_1_pipeline.png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # Fig 5.2 - Microservices
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5.5); ax.axis("off")
    services = [
        (0.3, 4.5, "ingestion\n:8001"), (2.0, 4.5, "normalization\n:8002"), (3.7, 4.5, "enrichment\n:8003"),
        (5.4, 4.5, "graph_builder\n:8004"), (7.1, 4.5, "detection\n:8005"), (8.6, 4.5, "case_mgmt\n:8006"),
        (1.5, 2.8, "llm_triage\n:8007"), (3.5, 2.8, "response\n:8008"), (5.5, 2.8, "approval\n:8009"),
        (7.5, 2.8, "audit\n:8010"), (4.5, 1.2, "frontend_gateway\n:8080"), (4.5, 0.1, "React UI\n:3000"),
    ]
    stores = [(0.5, 1.2, "Redpanda\n(Kafka)"), (2.5, 1.2, "Neo4j"), (6.5, 1.2, "Postgres"), (8.5, 1.2, "OpenSearch")]
    for x, y, lbl in services:
        ax.add_patch(FancyBboxPatch((x, y), 1.4, 0.85, boxstyle="round,pad=0.04", fc="#dbeafe", ec="#1d4ed8", lw=1.2))
        ax.text(x + 0.7, y + 0.42, lbl, ha="center", va="center", fontsize=6.5, fontweight="bold")
    for x, y, lbl in stores:
        ax.add_patch(FancyBboxPatch((x, y), 1.4, 0.75, boxstyle="round,pad=0.04", fc="#f0fdf4", ec="#16a34a", lw=1.2))
        ax.text(x + 0.7, y + 0.38, lbl, ha="center", va="center", fontsize=6.5, fontweight="bold")
    ax.set_title("Fig. 5.2 — Microservices Architecture", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(DIAGRAMS / "fig_5_2_microservices.png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # Fig 5.3 - Sequence
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")
    actors = ["Telemetry", "Detection", "Case Mgmt", "LLM Triage", "Analyst", "Approval", "Audit"]
    xs = [0.8 + i * 1.25 for i in range(len(actors))]
    for i, (x, a) in enumerate(zip(xs, actors)):
        ax.plot([x, x], [0.5, 4.5], "k--", lw=0.8, alpha=0.4)
        ax.text(x, 4.7, a, ha="center", fontsize=7, fontweight="bold")
    steps = [(0,1,4.2,"Alert"),(1,2,3.8,"Create case"),(2,3,3.4,"Evidence request"),(3,2,3.0,"Triage report"),
             (2,4,2.6,"Present case"),(4,5,2.2,"Approve action"),(5,6,1.8,"Log decision")]
    for s, d, y, lbl in steps:
        ax.annotate("", xy=(xs[d], y), xytext=(xs[s], y), arrowprops=dict(arrowstyle="->", color="#2563eb", lw=1.2))
        ax.text((xs[s]+xs[d])/2, y+0.08, lbl, ha="center", fontsize=6.5)
    ax.set_title("Fig. 5.3 — Alert-to-Response Sequence", fontsize=11, fontweight="bold")
    fig.tight_layout()
    fig.savefig(DIAGRAMS / "fig_5_3_sequence.png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build():
    _make_diagrams()

    src_doc = Document(str(SRC))
    body = src_doc.element.body

    # Keep cover + college certificate only (paragraphs 0..48 and signature table 0)
    paras = body.findall(qn("w:p"))
    for p in paras[49:]:
        body.remove(p)

    # Remove chapter tables (keep only certificate table at index 0)
    tables = body.findall(qn("w:tbl"))
    for tbl in tables[1:]:
        body.remove(tbl)

    doc = src_doc
    # Remove inline section breaks inside paragraphs; keep one trailing sectPr for section 0
    for p in body.findall(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            sp = pPr.find(qn("w:sectPr"))
            if sp is not None:
                pPr.remove(sp)

    _clear_header_footer(doc.sections[0])
    _set_margins(doc.sections[0])

    # --- Section 2: Front matter (roman numerals) ---
    doc.add_section(WD_SECTION.NEW_PAGE)
    sec_front = doc.sections[-1]
    _set_margins(sec_front)
    _setup_footer(sec_front, roman=True)

    # Acknowledgement
    _para(doc, "ACKNOWLEDGEMENT", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    for t in [
        "I express my sincere gratitude to Charotar University of Science and Technology (CHARUSAT) and the Department of Computer Engineering, DEPSTAR, for providing me the opportunity to undertake Summer Internship-II (CE446) as a project-based learning experience.",
        "I am deeply thankful to my internal guide Dr. Parth Goel, Assistant Professor, Department of Computer Science & Engineering, DEPSTAR, CHARUSAT, for his continuous guidance, constructive feedback, and academic support throughout the development of AegisSOC.",
        "I would also like to thank Digiesale, Vadodara, and especially Mr. Nikunj Bhavsar (HR) for facilitating the internship environment and encouraging practical exposure to enterprise-grade software engineering practices.",
        "Finally, I thank my family, friends, and classmates for their motivation and support during this internship.",
    ]:
        _para(doc, t, before=6)
    _para(doc, "", before=24)
    _para(doc, "Patel Harshil Dharmendra", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, "23DCE081", align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Abstract
    doc.add_page_break()
    _para(doc, "ABSTRACT", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    for t in [
        "Modern Security Operations Centers (SOCs) process millions of telemetry events daily, yet analysts must manually correlate fragmented alerts into coherent attack narratives. This internship project addresses that gap by designing and implementing AegisSOC, an AI-assisted SOC triage platform that ingests multi-source security telemetry, builds a temporal entity graph, detects suspicious activity using deterministic rules and graph analytics, and produces evidence-grounded triage reports for human analysts.",
        "The scope of the internship covered the complete software lifecycle: requirement analysis, microservices architecture design, backend service implementation, React analyst dashboard development, Docker-based deployment, and evaluation using labeled demo scenarios. The platform comprises eleven backend services, a frontend gateway with JWT authentication and role-based access control, and supporting infrastructure including Kafka (Redpanda), Neo4j, PostgreSQL, Redis, and OpenSearch.",
        "Actual work performed included canonical event schema design, Sigma-like detection rule integration, graph-based entity correlation, LLM-assisted (but not LLM-driven) triage summarization with citation grounding, human-in-the-loop response approval workflows, and end-to-end demo scenarios covering phishing-to-ransomware chains, false-positive suppression, and repeat-attacker infrastructure linking.",
        "The conclusion of this work is that a graph-aware, microservices SOC platform can materially improve analyst situational awareness while keeping detection explainable and keeping disruptive response actions under explicit human approval. The implemented system demonstrates production-shaped architecture at internship scale and provides a clear roadmap for horizontal scaling to enterprise telemetry volumes.",
    ]:
        _para(doc, t, before=6)

    # TOC
    doc.add_page_break()
    _para(doc, "TABLE OF CONTENTS", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    for title, page, is_chapter in TOC_ENTRIES:
        _toc_line(doc, title, page, bold=is_chapter and title.startswith("CHAPTER") or title in ("CONCLUSION", "REFERENCES", "ACKNOWLEDGEMENT", "ABSTRACT", "LIST OF FIGURES", "LIST OF TABLES", "DESCRIPTION OF COMPANY / ORGANIZATION"))

    # LOF
    doc.add_page_break()
    _para(doc, "LIST OF FIGURES", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    for num, title, page in LOF:
        _toc_line(doc, f"Fig. {num}   {title}", page)

    # LOT
    doc.add_page_break()
    _para(doc, "LIST OF TABLES", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    for num, title, page in LOT:
        _toc_line(doc, f"Table {num}   {title}", page)

    # Company description
    doc.add_page_break()
    _para(doc, "DESCRIPTION OF COMPANY / ORGANIZATION", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    for t in [
        "Digiesale is a technology-oriented organization based in Vadodara, Gujarat, focused on digital solutions, software services, and practical industry exposure for engineering interns. The company operates in the regional IT ecosystem and collaborates with academic institutions to bridge classroom learning with real-world project execution.",
        "Location and structure: Digiesale is located in Vadodara, a major industrial and educational hub in Gujarat. The organization follows a compact administrative structure suitable for agile software delivery, with functional coordination across human resources, technical mentoring, and project supervision. During this internship, the candidate worked in a project-based mode under organizational supervision while maintaining academic guidance from CHARUSAT.",
        "Main functions: The company's internship framework emphasizes end-to-end software development, including requirement understanding, architecture design, implementation, testing, documentation, and deployment. For this project, the focus was on building a security operations platform aligned with modern SOC workflows and cloud-native engineering practices.",
        "Customer profile and relevance: Digiesale serves clients seeking digital transformation and custom software solutions. The AegisSOC project aligns with the growing market demand for security analytics, SIEM augmentation, and AI-assisted incident response tools that reduce analyst fatigue while preserving auditability and human oversight.",
    ]:
        _para(doc, t, before=6)

    # --- Section 3: Chapters (arabic page numbers) ---
    doc.add_section(WD_SECTION.NEW_PAGE)
    sec_body = doc.sections[-1]
    _set_margins(sec_body)
    _setup_footer(sec_body, roman=False)

    # CHAPTER 1
    _chapter(doc, "CHAPTER 1  PROJECT DEFINITION")
    _section(doc, "1.1", "Introduction")
    _para(doc, "Security Operations Centers (SOCs) are responsible for detecting, investigating, and responding to cyber threats across enterprise environments. As log volumes grow into hundreds of millions of events per day, analysts struggle to separate true attacks from noise. Traditional SIEM platforms generate alerts but rarely reconstruct full attack paths automatically.")
    _para(doc, "AegisSOC is defined as an AI-assisted SOC triage platform whose primary objective is to transform raw telemetry into readable attack stories. The platform follows a strict design principle: the Large Language Model (LLM) is never the primary detector. Detection is performed by deterministic Sigma-like rules, correlation logic, graph feature extraction, and a calibrated machine learning ensemble. The LLM is invoked only after an alert or case exists, and only to summarize evidence, map MITRE ATT&CK techniques, and draft analyst-facing recommendations.", before=6)
    _section(doc, "1.2", "Problem Statement")
    _para(doc, "Existing SOC workflows suffer from three recurring problems: (1) alert fatigue caused by high false-positive rates, (2) loss of context because events remain siloed across endpoints, network, identity, and cloud sources, and (3) slow investigation because analysts manually pivot between tools to reconstruct timelines and entity relationships.")
    _section(doc, "1.3", "Project Objectives")
    for o in [
        "Design a microservices pipeline for ingest → normalize → enrich → graph → detect → case → triage → response.",
        "Build a temporal entity graph linking users, hosts, processes, files, IPs, domains, and techniques.",
        "Implement explainable detection using rules, correlation, and graph analytics.",
        "Provide an analyst dashboard for alert review, investigation, attack-path visualization, and approval workflows.",
        "Demonstrate end-to-end behavior using labeled synthetic scenarios with ground-truth attack labels.",
    ]:
        _bullet(doc, o)
    _section(doc, "1.4", "Scope")
    _para(doc, "The project scope includes eleven backend services, a React analyst console, Docker Compose deployment, sample datasets (2,880+ synthetic events), twenty Sigma-like rules, three demo scenarios, evaluation scripts, and architecture/security documentation. Out of scope for this internship phase are GPU-trained graph neural networks, live external threat-intel connectors, and multi-tenant physical isolation.")

    # CHAPTER 2
    _chapter(doc, "CHAPTER 2  DESCRIPTION")
    _section(doc, "2.1", "System Overview")
    _para(doc, "AegisSOC follows an event-driven microservices architecture. Telemetry enters through the ingestion service, is normalized into a canonical schema, enriched with ATT&CK tags and threat-intel matches, written into a Neo4j graph, analyzed by the detection engine, clustered into cases, summarized by the LLM triage service, and optionally routed to response playbooks requiring human approval for disruptive actions.")
    _section(doc, "2.2", "Architecture Layers")
    _subsection(doc, "2.2.1", "Data ingestion and streaming layer")
    _para(doc, "The ingestion service accepts raw events from Sysmon, Zeek, Suricata, EDR, Active Directory, CloudTrail, Kubernetes audit logs, and email/phishing sources. Valid events are published to Kafka topics; malformed records are sent to a dead-letter queue (DLQ). Redpanda provides the Kafka-compatible streaming backbone.")
    _subsection(doc, "2.2.2", "Processing and enrichment layer")
    _para(doc, "Normalization maps source-specific fields to CanonicalEvent objects with UTC timestamps and confidence metadata. Enrichment adds asset criticality, identity resolution, MITRE ATT&CK technique tags, and local threat-intel matches cached in Redis.")
    _subsection(doc, "2.2.3", "Graph and detection layer")
    _para(doc, "The graph builder performs entity resolution and upserts typed nodes and edges (e.g., Process spawned Process, Process created File, IP resolved_to Domain). Detection combines Sigma-like rules, temporal correlation, graph features (degree, rare-edge score, path-to-known-bad), and an ensemble scorer. Results are indexed in OpenSearch and emitted as alerts.")
    _subsection(doc, "2.2.4", "Analyst workflow layer")
    _para(doc, "Case management clusters related alerts, maintains timelines, and stores analyst feedback. The frontend gateway exposes a unified JWT-protected API to a React dashboard. LLM triage generates grounded reports with citations. Response policy recommends actions; disruptive actions require approval through the approval service and are logged in the append-only audit trail.")
    _section(doc, "2.3", "Technology Stack")
    _para(doc, "Backend services are implemented in Python 3.11 with FastAPI. The frontend uses React, TypeScript, and Vite. Infrastructure components include Docker Compose, PostgreSQL, Neo4j, Redis, OpenSearch, Redpanda, Prometheus, and optional Grafana/OpenTelemetry profiles. Kubernetes manifests and Helm charts are included for production-shaped deployment.")

    # CHAPTER 3
    _chapter(doc, "CHAPTER 3  SOFTWARE AND HARDWARE REQUIREMENTS")
    _section(doc, "3.1", "Software Requirements")
    _add_table(doc, ["Component", "Requirement"], [
        ["Operating System", "Linux/macOS/Windows with Docker support"],
        ["Docker & Docker Compose", "v2.x for full-stack deployment"],
        ["Python", "3.11+ for backend services and scripts"],
        ["Node.js & npm", "18+ for React frontend development"],
        ["Git", "Version control and collaboration"],
        ["Web Browser", "Chrome/Firefox/Edge for analyst console"],
        ["Optional LLM API", "OpenAI-compatible key for live triage calls"],
    ], "Table 3.1   Software Requirements")
    _section(doc, "3.2", "Hardware Requirements")
    _add_table(doc, ["Component", "Minimum", "Recommended"], [
        ["Processor", "4-core CPU", "8-core CPU or Apple Silicon equivalent"],
        ["RAM", "8 GB", "16 GB or higher"],
        ["Storage", "20 GB free SSD", "40 GB free SSD"],
        ["Network", "Localhost access for services", "Stable internet for LLM API (optional)"],
    ], "Table 3.2   Hardware Requirements")

    # CHAPTER 4
    _chapter(doc, "CHAPTER 4  MAJOR FUNCTIONALITY")
    _section(doc, "4.1", "Microservices")
    _add_table(doc, ["Service", "Port", "Responsibility"], [
        ["ingestion", "8001", "Raw telemetry intake, DLQ, replay API"],
        ["normalization", "8002", "Canonical event schema mapping"],
        ["enrichment", "8003", "ATT&CK tagging, intel matching"],
        ["graph_builder", "8004", "Temporal entity graph writes (Neo4j)"],
        ["detection", "8005", "Rules, correlation, ensemble scoring"],
        ["case_management", "8006", "Cases, timelines, analyst feedback"],
        ["llm_triage", "8007", "Evidence-grounded triage reports"],
        ["response_policy", "8008", "Playbook/action recommendation"],
        ["approval", "8009", "Human-in-the-loop approval gate"],
        ["audit", "8010", "Append-only audit trail"],
        ["frontend_gateway", "8080", "Public API, JWT auth, RBAC"],
    ], "Table 4.1   Microservices and Responsibilities")
    _section(doc, "4.2", "Key Features")
    for f in [
        "Multi-source telemetry ingestion with scenario replay for demos and evaluation.",
        "Temporal entity graph with neighborhood queries and attack-path analysis.",
        "Sigma-like rule engine with correlation and graph-enhanced risk scoring.",
        "Case clustering, investigation workspace, timeline, and evidence panels.",
        "LLM triage reports with groundedness validation and ATT&CK mapping.",
        "Response recommendations with approval workflow for disruptive actions.",
        "Prometheus metrics, health checks, and optional Grafana dashboards.",
    ]:
        _bullet(doc, f)
    _section(doc, "4.3", "Demo Scenarios")
    _add_table(doc, ["Scenario", "Severity", "Purpose"], [
        ["phishing_ransomware_chain", "Critical", "Full kill chain: phishing → ransomware"],
        ["benign_admin_false_positive", "Low", "False-positive suppression demonstration"],
        ["repeat_attacker_infra", "High", "Graph memory across historical incidents"],
    ], "Table 4.2   Demo Scenarios and Expected Outcomes")

    # CHAPTER 5
    _chapter(doc, "CHAPTER 5  FLOW CHART")
    _section(doc, "5.1", "End-to-End Pipeline Flow")
    _para(doc, "The following diagram illustrates the primary data flow from telemetry sources to analyst-facing outputs. Events move through ingestion and normalization, receive enrichment and graph updates, trigger detection, create or update cases, and optionally invoke LLM triage and response workflows.")
    _pic(doc, DIAGRAMS / "fig_5_1_pipeline.png")
    _caption(doc, "Fig. 5.1   AegisSOC End-to-End Data Flow")
    _section(doc, "5.2", "Microservices Architecture")
    _para(doc, "Figure 5.2 shows the eleven backend microservices, frontend gateway, React analyst UI, and supporting data stores (Redpanda, Neo4j, PostgreSQL, OpenSearch) that comprise the AegisSOC deployment topology.")
    _pic(doc, DIAGRAMS / "fig_5_2_microservices.png")
    _caption(doc, "Fig. 5.2   Microservices Architecture Diagram")
    _section(doc, "5.3", "Alert-to-Response Sequence")
    _para(doc, "When detection emits an alert, case management clusters it with related alerts. An analyst opens the Investigation Workspace, reviews the attack-path graph and timeline, requests an LLM triage report, evaluates recommended actions, and submits disruptive actions for senior/admin approval. All decisions are recorded in the audit service.")
    _pic(doc, DIAGRAMS / "fig_5_3_sequence.png")
    _caption(doc, "Fig. 5.3   Alert-to-Response Sequence Diagram")

    # CHAPTER 6
    _chapter(doc, "CHAPTER 6  SCREENSHOTS OF PROJECT OUTPUT")
    for num, title, desc, img in SCREENSHOTS:
        sec_num = num.split(".")[1]
        _section(doc, f"6.{sec_num}", title)
        _para(doc, desc)
        _pic(doc, ASSETS / img, width=Inches(5.6))
        _caption(doc, f"Fig. {num}   {title}")

    # CHAPTER 7-9, Conclusion, References
    _chapter(doc, "CHAPTER 7  LIMITATIONS OF PROJECT")
    _para(doc, "Although AegisSOC implements production-shaped architecture, several limitations remain appropriate for an internship-scale deployment:")
    for lim in [
        "Graph ML uses a lightweight GraphSAGE-style numpy scorer rather than a GPU-trained heterogeneous GNN.",
        "Live VirusTotal/OpenCTI connectors are optional stubs; local JSON intel feeds are used by default.",
        "Prompt-injection defenses are heuristic pattern filters rather than a dedicated classifier.",
        "Tenant isolation is logical (tenant_id scoping) rather than physically isolated infrastructure.",
        "Online reinforcement learning is not enabled; response ranking uses offline-evaluated contextual bandits.",
        "Local Docker Compose deployment does not demonstrate 100M logs/day throughput without horizontal scaling.",
    ]:
        _bullet(doc, lim)

    _chapter(doc, "CHAPTER 8  OUTCOME")
    _para(doc, "The AegisSOC internship project successfully delivered a working SOC triage platform with the following outcomes:")
    for o in [
        "Eleven backend microservices and a React analyst dashboard running in Docker Compose.",
        "End-to-end demo scenarios covering ransomware chains, false positives, and repeat-attacker graph memory.",
        "Evidence-grounded LLM triage integrated without compromising deterministic detection.",
        "Human-in-the-loop approval workflow for disruptive response actions with audit logging.",
        "Architecture decision records, threat model, evaluation scripts, and deployment manifests.",
        "Practical experience in security domain modeling, streaming pipelines, and full-stack integration.",
    ]:
        _bullet(doc, o)
    _para(doc, "Evaluation scripts measure detection precision/recall, LLM groundedness, and ingestion throughput, providing quantitative evidence of platform behavior beyond manual demo observation.", before=6)

    _chapter(doc, "CHAPTER 9  FUTURE ENHANCEMENT")
    for e in [
        "Scale Redpanda to a multi-broker cluster with KEDA lag-based autoscaling.",
        "Migrate Neo4j Community to Causal Cluster or AuraDB with read replicas.",
        "Integrate live threat-intel feeds (OpenCTI, MISP, VirusTotal) with caching and rate limits.",
        "Train and deploy a heterogeneous GNN for improved graph-based risk scoring.",
        "Harden ingestion edge authentication with mTLS and per-source API keys.",
        "Add SOAR connector integrations for approved response actions beyond dry-run mode.",
        "Expand RBAC roles (viewer, tenant_admin, playbook_editor) and fine-grained audit search.",
    ]:
        _bullet(doc, e)

    _chapter(doc, "CONCLUSION")
    _para(doc, "This internship project resulted in the design and implementation of AegisSOC, a comprehensive AI-assisted SOC triage platform that addresses alert fatigue and investigation latency through graph-aware correlation and evidence-grounded automation. By keeping detection deterministic and restricting the LLM to analyst copilot responsibilities, the system maintains explainability and operational trust.")
    _para(doc, "The completed platform demonstrates a realistic path from telemetry ingestion to human-approved response, supported by documented architecture, security controls, and scalability plans. The experience strengthened skills in microservices engineering, cybersecurity domain modeling, and full-stack delivery — aligning academic learning at CHARUSAT with practical software development exposure at Digiesale.", before=6)

    _chapter(doc, "REFERENCES")
    refs = [
        '[1] MITRE Corporation, "MITRE ATT&CK Framework," https://attack.mitre.org/, accessed July 2026.',
        '[2] Sigma Project, "Sigma Generic Signature Format," https://github.com/SigmaHQ/sigma, accessed July 2026.',
        '[3] FastAPI Documentation, "FastAPI framework, high performance, easy to learn," https://fastapi.tiangolo.com/, accessed July 2026.',
        '[4] Neo4j Inc., "Neo4j Graph Database Documentation," https://neo4j.com/docs/, accessed July 2026.',
        '[5] Apache Software Foundation, "Apache Kafka Documentation," https://kafka.apache.org/documentation/, accessed July 2026.',
        '[6] OpenSearch Project, "OpenSearch Documentation," https://opensearch.org/docs/latest/, accessed July 2026.',
        '[7] React Team, "React Documentation," https://react.dev/, accessed July 2026.',
        '[8] Docker Inc., "Docker Documentation," https://docs.docker.com/, accessed July 2026.',
        '[9] NIST, "NIST Cybersecurity Framework," https://www.nist.gov/cyberframework, accessed July 2026.',
        '[10] Redpanda Data, "Redpanda Documentation," https://docs.redpanda.com/, accessed July 2026.',
    ]
    for r in refs:
        _para(doc, r, spacing=1.0, after=4)

    # Set default font
    try:
        doc.styles["Normal"].font.name = FONT
        doc.styles["Normal"].font.size = Pt(12)
    except Exception:
        pass

    doc.save(str(OUT))
    print(f"Saved formatted report: {OUT}")


if __name__ == "__main__":
    build()
