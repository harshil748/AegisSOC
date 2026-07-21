#!/usr/bin/env python3
"""Generate AegisSOC Summer Internship-II project report (report.docx)."""

from __future__ import annotations

import os
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docxcompose.composer import Composer

ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = Path("/Users/harshil/Downloads")
OUT = ROOT / "report.docx"
OUT_DL = DOWNLOADS / "report.docx"
ASSETS = ROOT / "docs" / "report-assets"
ASSETS.mkdir(parents=True, exist_ok=True)

STUDENT_ID = "23DCE081"
FONT = "Times New Roman"


def _set_run_font(run, size=12, bold=False, caps=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    if caps:
        run.font.all_caps = True


def _fmt_paragraph(p, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.5, before=0, after=6):
    pf = p.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Pt(0)
    for run in p.runs:
        _set_run_font(run, size=size, bold=bold)


def _add_paragraph(doc, text, *, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, style=None, spacing=1.5, before=0, after=6):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    _fmt_paragraph(p, size=size, bold=bold, align=align, spacing=spacing, before=before, after=after)
    return p


def _add_chapter(doc, title: str):
    doc.add_page_break()
    p = _add_paragraph(doc, title, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=12)
    for run in p.runs:
        run.font.all_caps = True
    return p


def _add_section(doc, number: str, title: str):
    p = _add_paragraph(doc, f"{number}   {title.upper()}", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)
    return p


def _add_subsection(doc, number: str, title: str):
    t = title[0].upper() + title[1:] if title else title
    return _add_paragraph(doc, f"{number}   {t}", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=6)


def _add_bullet(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"• {text}")
    _set_run_font(run, size=12)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.0
    pf.left_indent = Inches(0.35)
    pf.space_after = Pt(3)
    return p


def _add_caption(doc, text: str, *, center=True):
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    return _add_paragraph(doc, text, size=12, align=align, spacing=1.0, before=6, after=12)


def _add_picture(doc, path: Path, width=Inches(5.8)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _create_field_run(field_code: str):
    run = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and Update Field"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run.append(fld_begin)
    run.append(instr)
    run.append(fld_sep)
    run.append(OxmlElement("w:r"))
    run[-1].append(fld_text)
    run.append(fld_end)
    return run


def _add_toc(doc):
    _add_paragraph(doc, "TABLE OF CONTENTS", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    p = doc.add_paragraph()
    p._p.append(_create_field_run(' TOC \\o "1-3" \\h \\z \\u '))
    _add_paragraph(doc, "", size=12)


def _set_margins(section, left=1.25, right=1.0, top=1.0, bottom=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def _header_footer(section, chapter_title: str = "", roman: bool = False):
    header = section.header
    header.is_linked_to_previous = False
    hp_left = header.paragraphs[0]
    hp_left.text = STUDENT_ID
    _fmt_paragraph(hp_left, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0)
    hp_right = header.add_paragraph(chapter_title)
    _fmt_paragraph(hp_right, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, spacing=1.0)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp1 = footer.paragraphs[0]
    fp1.text = "DEPSTAR"
    _fmt_paragraph(fp1, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0)
    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp2.add_run()
    if roman:
        run.text = "i"
    else:
        run.text = "1"
    _set_run_font(run, size=10)
    fp3 = footer.add_paragraph("Department of Computer Engineering")
    _fmt_paragraph(fp3, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, spacing=1.0)


def _add_table(doc, headers, rows, caption: str):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                _set_run_font(r, size=12, bold=True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    _set_run_font(r, size=12)
    _add_caption(doc, caption)


def _make_architecture_diagram(path: Path):
    try:
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    except ImportError:
        return False

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")

    boxes = [
        (0.3, 4.6, "Telemetry\nSources"),
        (2.0, 4.6, "Ingestion"),
        (3.7, 4.6, "Normalization"),
        (5.4, 4.6, "Enrichment"),
        (7.1, 4.6, "Graph Builder"),
        (8.6, 4.6, "Neo4j"),
        (5.4, 2.8, "Detection"),
        (3.7, 1.2, "Case Mgmt"),
        (5.4, 1.2, "LLM Triage"),
        (7.1, 1.2, "Response"),
        (8.6, 1.2, "Analyst UI"),
    ]
    coords = {}
    for x, y, label in boxes:
        patch = FancyBboxPatch((x, y), 1.4, 0.9, boxstyle="round,pad=0.05", linewidth=1.2, edgecolor="#2563eb", facecolor="#eff6ff")
        ax.add_patch(patch)
        ax.text(x + 0.7, y + 0.45, label, ha="center", va="center", fontsize=8, fontweight="bold")
        coords[label.split("\n")[0]] = (x + 1.4, y + 0.45)

    def arrow(a, b):
        ax.add_patch(FancyArrowPatch(a, b, arrowstyle="->", mutation_scale=12, linewidth=1.1, color="#334155"))

    flow = [
        ((0.3 + 1.4, 5.05), (2.0, 5.05)),
        ((2.0 + 1.4, 5.05), (3.7, 5.05)),
        ((3.7 + 1.4, 5.05), (5.4, 5.05)),
        ((5.4 + 1.4, 5.05), (7.1, 5.05)),
        ((7.1 + 1.4, 5.05), (8.6, 5.05)),
        ((5.4 + 0.7, 4.6), (5.4 + 0.7, 3.7)),
        ((5.4 + 0.7, 2.8), (4.4, 2.1)),
        ((4.4, 1.2 + 0.9), (4.4, 2.1)),
        ((3.7 + 1.4, 1.65), (5.4, 1.65)),
        ((5.4 + 1.4, 1.65), (7.1, 1.65)),
        ((7.1 + 1.4, 1.65), (8.6, 1.65)),
    ]
    for s, e in flow:
        arrow(s, e)

    ax.set_title("AegisSOC End-to-End Data Flow", fontsize=12, fontweight="bold", pad=12)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return True


def _company_certificate(doc: Document):
    doc.add_page_break()
    _add_paragraph(doc, "COMPANY CERTIFICATE", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=24)
    _add_paragraph(
        doc,
        'This is to certify that Mr. Patel Harshil Dharmendra (23DCE081) has successfully completed '
        'the Summer Internship-II (Project Based) titled "AegisSOC: AI-Assisted SOC Triage Platform" '
        "at Digiesale, Vadodara, Gujarat, during the period May 2026 to July 2026. During the internship, "
        "the candidate demonstrated sincerity, technical competence, and professional discipline while "
        "designing and implementing a microservices-based security operations platform.",
        size=12,
    )
    _add_paragraph(
        doc,
        "The work carried out by the candidate is original to the best of our knowledge and fulfills "
        "the organizational requirements of the internship program.",
        size=12,
        before=12,
    )
    _add_paragraph(doc, "", size=12, before=36)
    _add_paragraph(doc, "For Digiesale", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, "Vadodara, Gujarat", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, "", size=12, before=48)
    _add_paragraph(doc, "____________________________", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, "Nikunj Bhavsar", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, "HR Manager", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, "Digiesale, Vadodara", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, "Date: July 2026", size=12, align=WD_ALIGN_PARAGRAPH.LEFT, before=12)


def _acknowledgement(doc: Document):
    doc.add_page_break()
    _add_paragraph(doc, "ACKNOWLEDGEMENT", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    _add_paragraph(
        doc,
        "I express my sincere gratitude to Charotar University of Science and Technology (CHARUSAT) and "
        "the Department of Computer Engineering, DEPSTAR, for providing me the opportunity to undertake "
        "Summer Internship-II (CE446) as a project-based learning experience.",
        size=12,
    )
    _add_paragraph(
        doc,
        "I am deeply thankful to my internal guide Dr. Parth Goel, Assistant Professor, Department of "
        "Computer Science & Engineering, DEPSTAR, CHARUSAT, for his continuous guidance, constructive "
        "feedback, and academic support throughout the development of AegisSOC.",
        size=12,
        before=12,
    )
    _add_paragraph(
        doc,
        "I would also like to thank Digiesale, Vadodara, and especially Mr. Nikunj Bhavsar (HR) for "
        "facilitating the internship environment and encouraging practical exposure to enterprise-grade "
        "software engineering practices.",
        size=12,
        before=12,
    )
    _add_paragraph(
        doc,
        "Finally, I thank my family, friends, and classmates for their motivation and support during this internship.",
        size=12,
        before=12,
    )
    _add_paragraph(doc, "", size=12, before=36)
    _add_paragraph(doc, "Patel Harshil Dharmendra", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, "23DCE081", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)


def _abstract(doc: Document):
    doc.add_page_break()
    _add_paragraph(doc, "ABSTRACT", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    _add_paragraph(
        doc,
        "Modern Security Operations Centers (SOCs) process millions of telemetry events daily, yet analysts "
        "must manually correlate fragmented alerts into coherent attack narratives. This internship project "
        "addresses that gap by designing and implementing AegisSOC, an AI-assisted SOC triage platform that "
        "ingests multi-source security telemetry, builds a temporal entity graph, detects suspicious activity "
        "using deterministic rules and graph analytics, and produces evidence-grounded triage reports for human analysts.",
        size=12,
    )
    _add_paragraph(
        doc,
        "The scope of the internship covered the complete software lifecycle: requirement analysis, microservices "
        "architecture design, backend service implementation, React analyst dashboard development, Docker-based "
        "deployment, and evaluation using labeled demo scenarios. The platform comprises eleven backend services, "
        "a frontend gateway with JWT authentication and role-based access control, and supporting infrastructure "
        "including Kafka (Redpanda), Neo4j, PostgreSQL, Redis, and OpenSearch.",
        size=12,
        before=12,
    )
    _add_paragraph(
        doc,
        "Actual work performed included canonical event schema design, Sigma-like detection rule integration, "
        "graph-based entity correlation, LLM-assisted (but not LLM-driven) triage summarization with citation "
        "grounding, human-in-the-loop response approval workflows, and end-to-end demo scenarios covering "
        "phishing-to-ransomware chains, false-positive suppression, and repeat-attacker infrastructure linking.",
        size=12,
        before=12,
    )
    _add_paragraph(
        doc,
        "The conclusion of this work is that a graph-aware, microservices SOC platform can materially improve "
        "analyst situational awareness while keeping detection explainable and keeping disruptive response actions "
        "under explicit human approval. The implemented system demonstrates production-shaped architecture at "
        "internship scale and provides a clear roadmap for horizontal scaling to enterprise telemetry volumes.",
        size=12,
        before=12,
    )


def _list_of_figures(doc: Document):
    doc.add_page_break()
    _add_paragraph(doc, "LIST OF FIGURES", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    items = [
        "Fig. 5.1   AegisSOC End-to-End Data Flow",
        "Fig. 6.1   Investigation Workspace with Triage Report",
        "Fig. 6.2   Replay and Demo Scenario Console",
    ]
    for item in items:
        _add_paragraph(doc, item, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, after=3)


def _list_of_tables(doc: Document):
    doc.add_page_break()
    _add_paragraph(doc, "LIST OF TABLES", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    items = [
        "Table 3.1   Software Requirements",
        "Table 3.2   Hardware Requirements",
        "Table 4.1   Microservices and Responsibilities",
        "Table 4.2   Demo Scenarios and Expected Outcomes",
    ]
    for item in items:
        _add_paragraph(doc, item, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, spacing=1.0, after=3)


def _company_description(doc: Document):
    doc.add_page_break()
    _add_paragraph(doc, "DESCRIPTION OF COMPANY / ORGANIZATION", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=18)
    _add_paragraph(
        doc,
        "Digiesale is a technology-oriented organization based in Vadodara, Gujarat, focused on digital "
        "solutions, software services, and practical industry exposure for engineering interns. The company "
        "operates in the regional IT ecosystem and collaborates with academic institutions to bridge classroom "
        "learning with real-world project execution.",
        size=12,
    )
    _add_paragraph(
        doc,
        "Location and structure: Digiesale is located in Vadodara, a major industrial and educational hub in "
        "Gujarat. The organization follows a compact administrative structure suitable for agile software "
        "delivery, with functional coordination across human resources, technical mentoring, and project "
        "supervision. During this internship, the candidate worked in a project-based mode under organizational "
        "supervision while maintaining academic guidance from CHARUSAT.",
        size=12,
        before=12,
    )
    _add_paragraph(
        doc,
        "Main functions: The company's internship framework emphasizes end-to-end software development, including "
        "requirement understanding, architecture design, implementation, testing, documentation, and deployment. "
        "For this project, the focus was on building a security operations platform aligned with modern SOC "
        "workflows and cloud-native engineering practices.",
        size=12,
        before=12,
    )
    _add_paragraph(
        doc,
        "Customer profile and relevance: Digiesale serves clients seeking digital transformation and custom "
        "software solutions. The AegisSOC project aligns with the growing market demand for security analytics, "
        "SIEM augmentation, and AI-assisted incident response tools that reduce analyst fatigue while preserving "
        "auditability and human oversight.",
        size=12,
        before=12,
    )


def _chapter1(doc: Document):
    _add_chapter(doc, "CHAPTER 1  PROJECT DEFINITION")
    _add_section(doc, "1.1", "Introduction")
    _add_paragraph(
        doc,
        "Security Operations Centers (SOCs) are responsible for detecting, investigating, and responding to "
        "cyber threats across enterprise environments. As log volumes grow into hundreds of millions of events "
        "per day, analysts struggle to separate true attacks from noise. Traditional SIEM platforms generate "
        "alerts but rarely reconstruct full attack paths automatically.",
        size=12,
    )
    _add_paragraph(
        doc,
        "AegisSOC is defined as an AI-assisted SOC triage platform whose primary objective is to transform raw "
        "telemetry into readable attack stories. The platform follows a strict design principle: the Large Language "
        "Model (LLM) is never the primary detector. Detection is performed by deterministic Sigma-like rules, "
        "correlation logic, graph feature extraction, and a calibrated machine learning ensemble. The LLM is invoked "
        "only after an alert or case exists, and only to summarize evidence, map MITRE ATT&CK techniques, and "
        "draft analyst-facing recommendations.",
        size=12,
        before=12,
    )
    _add_section(doc, "1.2", "Problem statement")
    _add_paragraph(
        doc,
        "Existing SOC workflows suffer from three recurring problems: (1) alert fatigue caused by high false-positive "
        "rates, (2) loss of context because events remain siloed across endpoints, network, identity, and cloud sources, "
        "and (3) slow investigation because analysts manually pivot between tools to reconstruct timelines and entity "
        "relationships.",
        size=12,
    )
    _add_section(doc, "1.3", "Project objectives")
    for obj in [
        "Design a microservices pipeline for ingest → normalize → enrich → graph → detect → case → triage → response.",
        "Build a temporal entity graph linking users, hosts, processes, files, IPs, domains, and techniques.",
        "Implement explainable detection using rules, correlation, and graph analytics.",
        "Provide an analyst dashboard for alert review, investigation, attack-path visualization, and approval workflows.",
        "Demonstrate end-to-end behavior using labeled synthetic scenarios with ground-truth attack labels.",
    ]:
        _add_bullet(doc, obj)
    _add_section(doc, "1.4", "Scope")
    _add_paragraph(
        doc,
        "The project scope includes eleven backend services, a React analyst console, Docker Compose deployment, "
        "sample datasets (2,880+ synthetic events), twenty Sigma-like rules, three demo scenarios, evaluation "
        "scripts, and architecture/security documentation. Out of scope for this internship phase are GPU-trained "
        "graph neural networks, live external threat-intel connectors, and multi-tenant physical isolation.",
        size=12,
    )


def _chapter2(doc: Document):
    _add_chapter(doc, "CHAPTER 2  DESCRIPTION")
    _add_section(doc, "2.1", "System overview")
    _add_paragraph(
        doc,
        "AegisSOC follows an event-driven microservices architecture. Telemetry enters through the ingestion "
        "service, is normalized into a canonical schema, enriched with ATT&CK tags and threat-intel matches, "
        "written into a Neo4j graph, analyzed by the detection engine, clustered into cases, summarized by the "
        "LLM triage service, and optionally routed to response playbooks requiring human approval for disruptive actions.",
        size=12,
    )
    _add_section(doc, "2.2", "Architecture layers")
    _add_subsection(doc, "2.2.1", "Data ingestion and streaming layer")
    _add_paragraph(
        doc,
        "The ingestion service accepts raw events from Sysmon, Zeek, Suricata, EDR, Active Directory, CloudTrail, "
        "Kubernetes audit logs, and email/phishing sources. Valid events are published to Kafka topics; malformed "
        "records are sent to a dead-letter queue (DLQ). Redpanda provides the Kafka-compatible streaming backbone.",
        size=12,
    )
    _add_subsection(doc, "2.2.2", "Processing and enrichment layer")
    _add_paragraph(
        doc,
        "Normalization maps source-specific fields to CanonicalEvent objects with UTC timestamps and confidence "
        "metadata. Enrichment adds asset criticality, identity resolution, MITRE ATT&CK technique tags, and local "
        "threat-intel matches cached in Redis.",
        size=12,
    )
    _add_subsection(doc, "2.2.3", "Graph and detection layer")
    _add_paragraph(
        doc,
        "The graph builder performs entity resolution and upserts typed nodes and edges (e.g., Process spawned Process, "
        "Process created File, IP resolved_to Domain). Detection combines Sigma-like rules, temporal correlation, "
        "graph features (degree, rare-edge score, path-to-known-bad), and an ensemble scorer. Results are indexed "
        "in OpenSearch and emitted as alerts.",
        size=12,
    )
    _add_subsection(doc, "2.2.4", "Analyst workflow layer")
    _add_paragraph(
        doc,
        "Case management clusters related alerts, maintains timelines, and stores analyst feedback. The frontend "
        "gateway exposes a unified JWT-protected API to a React dashboard. LLM triage generates grounded reports "
        "with citations. Response policy recommends actions; disruptive actions require approval through the approval "
        "service and are logged in the append-only audit trail.",
        size=12,
    )
    _add_section(doc, "2.3", "Technology stack")
    _add_paragraph(
        doc,
        "Backend services are implemented in Python 3.11 with FastAPI. The frontend uses React, TypeScript, and Vite. "
        "Infrastructure components include Docker Compose, PostgreSQL, Neo4j, Redis, OpenSearch, Redpanda, Prometheus, "
        "and optional Grafana/OpenTelemetry profiles. Kubernetes manifests and Helm charts are included for production-shaped deployment.",
        size=12,
    )


def _chapter3(doc: Document):
    _add_chapter(doc, "CHAPTER 3  SOFTWARE AND HARDWARE REQUIREMENTS")
    _add_section(doc, "3.1", "Software requirements")
    _add_table(
        doc,
        ["Component", "Requirement"],
        [
            ["Operating System", "Linux/macOS/Windows with Docker support"],
            ["Docker & Docker Compose", "v2.x for full-stack deployment"],
            ["Python", "3.11+ for backend services and scripts"],
            ["Node.js & npm", "18+ for React frontend development"],
            ["Git", "Version control and collaboration"],
            ["Web Browser", "Chrome/Firefox/Edge for analyst console"],
            ["Optional LLM API", "OpenAI-compatible key for live triage calls"],
        ],
        "Table 3.1   Software Requirements",
    )
    _add_section(doc, "3.2", "Hardware requirements")
    _add_table(
        doc,
        ["Component", "Minimum", "Recommended"],
        [
            ["Processor", "4-core CPU", "8-core CPU or Apple Silicon equivalent"],
            ["RAM", "8 GB", "16 GB or higher"],
            ["Storage", "20 GB free SSD", "40 GB free SSD"],
            ["Network", "Localhost access for services", "Stable internet for LLM API (optional)"],
        ],
        "Table 3.2   Hardware Requirements",
    )


def _chapter4(doc: Document):
    _add_chapter(doc, "CHAPTER 4  MAJOR FUNCTIONALITY")
    _add_section(doc, "4.1", "Microservices")
    _add_table(
        doc,
        ["Service", "Port", "Responsibility"],
        [
            ["ingestion", "8001", "Raw telemetry intake, DLQ, replay API"],
            ["normalization", "8002", "Canonical event schema mapping"],
            ["enrichment", "8003", "ATT&CK tagging, intel matching"],
            ["graph_builder", "8004", "Temporal entity graph writes (Neo4j)"],
            ["detection", "8005", "Rules, correlation, ensemble scoring"],
            ["case_management", "8006", "Cases, timelines, analyst feedback"],
            ["llm_triage", "8007", "Evidence-grounded triage reports"],
            ["response_policy", "8008", "Playbook/action recommendation"],
            ["approval", "8009", "Human-in-the-loop approval gate"],
            ["audit", "8010", "Append-only audit trail"],
            ["frontend_gateway", "8080", "Public API, JWT auth, RBAC"],
        ],
        "Table 4.1   Microservices and Responsibilities",
    )
    _add_section(doc, "4.2", "Key features")
    for feat in [
        "Multi-source telemetry ingestion with scenario replay for demos and evaluation.",
        "Temporal entity graph with neighborhood queries and attack-path analysis.",
        "Sigma-like rule engine with correlation and graph-enhanced risk scoring.",
        "Case clustering, investigation workspace, timeline, and evidence panels.",
        "LLM triage reports with groundedness validation and ATT&CK mapping.",
        "Response recommendations with approval workflow for disruptive actions.",
        "Prometheus metrics, health checks, and optional Grafana dashboards.",
    ]:
        _add_bullet(doc, feat)
    _add_section(doc, "4.3", "Demo scenarios")
    _add_table(
        doc,
        ["Scenario", "Severity", "Purpose"],
        [
            ["phishing_ransomware_chain", "Critical", "Full kill chain: phishing → ransomware"],
            ["benign_admin_false_positive", "Low", "False-positive suppression demonstration"],
            ["repeat_attacker_infra", "High", "Graph memory across historical incidents"],
        ],
        "Table 4.2   Demo Scenarios and Expected Outcomes",
    )


def _chapter5(doc: Document, diagram: Path):
    _add_chapter(doc, "CHAPTER 5  FLOW CHART")
    _add_section(doc, "5.1", "End-to-end pipeline flow")
    _add_paragraph(
        doc,
        "The following diagram illustrates the primary data flow from telemetry sources to analyst-facing outputs. "
        "Events move through ingestion and normalization, receive enrichment and graph updates, trigger detection, "
        "create or update cases, and optionally invoke LLM triage and response workflows.",
        size=12,
    )
    if diagram.exists():
        _add_picture(doc, diagram)
        _add_caption(doc, "Fig. 5.1   AegisSOC End-to-End Data Flow")
    _add_section(doc, "5.2", "Alert-to-response sequence")
    _add_paragraph(
        doc,
        "When detection emits an alert, case management clusters it with related alerts. An analyst opens the "
        "Investigation Workspace, reviews the attack-path graph and timeline, requests an LLM triage report, "
        "evaluates recommended actions, and submits disruptive actions for senior/admin approval. All decisions "
        "are recorded in the audit service.",
        size=12,
    )


def _chapter6(doc: Document, shots: dict[str, Path]):
    _add_chapter(doc, "CHAPTER 6  SCREENSHOTS OF PROJECT OUTPUT")
    _add_section(doc, "6.1", "Investigation workspace")
    _add_paragraph(
        doc,
        "The investigation workspace combines attack-path graph visualization, LLM triage report, "
        "timeline, and evidence panels for a selected case.",
        size=12,
    )
    if shots.get("investigate"):
        _add_picture(doc, shots["investigate"])
        _add_caption(doc, "Fig. 6.1   Investigation Workspace with Triage Report")
    _add_section(doc, "6.2", "Replay and demo console")
    _add_paragraph(
        doc,
        "The replay page executes canonical demo scenarios to validate detection, triage, and response end-to-end.",
        size=12,
    )
    if shots.get("replay"):
        _add_picture(doc, shots["replay"])
        _add_caption(doc, "Fig. 6.2   Replay and Demo Scenario Console")


def _chapter7(doc: Document):
    _add_chapter(doc, "CHAPTER 7  LIMITATIONS OF PROJECT")
    _add_paragraph(
        doc,
        "Although AegisSOC implements production-shaped architecture, several limitations remain appropriate for an "
        "internship-scale deployment:",
        size=12,
    )
    for lim in [
        "Graph ML uses a lightweight GraphSAGE-style numpy scorer rather than a GPU-trained heterogeneous GNN.",
        "Live VirusTotal/OpenCTI connectors are optional stubs; local JSON intel feeds are used by default.",
        "Prompt-injection defenses are heuristic pattern filters rather than a dedicated classifier.",
        "Tenant isolation is logical (tenant_id scoping) rather than physically isolated infrastructure.",
        "Online reinforcement learning is not enabled; response ranking uses offline-evaluated contextual bandits.",
        "Local Docker Compose deployment does not demonstrate 100M logs/day throughput without horizontal scaling.",
    ]:
        _add_bullet(doc, lim)


def _chapter8(doc: Document):
    _add_chapter(doc, "CHAPTER 8  OUTCOME")
    _add_paragraph(
        doc,
        "The AegisSOC internship project successfully delivered a working SOC triage platform with the following outcomes:",
        size=12,
    )
    for out in [
        "Eleven backend microservices and a React analyst dashboard running in Docker Compose.",
        "End-to-end demo scenarios covering ransomware chains, false positives, and repeat-attacker graph memory.",
        "Evidence-grounded LLM triage integrated without compromising deterministic detection.",
        "Human-in-the-loop approval workflow for disruptive response actions with audit logging.",
        "Architecture decision records, threat model, evaluation scripts, and deployment manifests.",
        "Practical experience in security domain modeling, streaming pipelines, and full-stack integration.",
    ]:
        _add_bullet(doc, out)
    _add_paragraph(
        doc,
        "Evaluation scripts measure detection precision/recall, LLM groundedness, and ingestion throughput, providing "
        "quantitative evidence of platform behavior beyond manual demo observation.",
        size=12,
        before=12,
    )


def _chapter9(doc: Document):
    _add_chapter(doc, "CHAPTER 9  FUTURE ENHANCEMENT")
    for enh in [
        "Scale Redpanda to a multi-broker cluster with KEDA lag-based autoscaling.",
        "Migrate Neo4j Community to Causal Cluster or AuraDB with read replicas.",
        "Integrate live threat-intel feeds (OpenCTI, MISP, VirusTotal) with caching and rate limits.",
        "Train and deploy a heterogeneous GNN for improved graph-based risk scoring.",
        "Harden ingestion edge authentication with mTLS and per-source API keys.",
        "Add SOAR connector integrations for approved response actions beyond dry-run mode.",
        "Expand RBAC roles (viewer, tenant_admin, playbook_editor) and fine-grained audit search.",
    ]:
        _add_bullet(doc, enh)


def _conclusion(doc: Document):
    _add_chapter(doc, "CONCLUSION")
    _add_paragraph(
        doc,
        "This internship project resulted in the design and implementation of AegisSOC, a comprehensive AI-assisted "
        "SOC triage platform that addresses alert fatigue and investigation latency through graph-aware correlation "
        "and evidence-grounded automation. By keeping detection deterministic and restricting the LLM to analyst "
        "copilot responsibilities, the system maintains explainability and operational trust.",
        size=12,
    )
    _add_paragraph(
        doc,
        "The completed platform demonstrates a realistic path from telemetry ingestion to human-approved response, "
        "supported by documented architecture, security controls, and scalability plans. The experience strengthened "
        "skills in microservices engineering, cybersecurity domain modeling, and full-stack delivery — aligning "
        "academic learning at CHARUSAT with practical software development exposure at Digiesale.",
        size=12,
        before=12,
    )


def _references(doc: Document):
    _add_chapter(doc, "REFERENCES")
    refs = [
        '[1] MITRE Corporation, "MITRE ATT&CK Framework," https://attack.mitre.org/, accessed July 2026.',
        '[2] Sigma Project, "Sigma Generic Signature Format," https://github.com/SigmaHQ/sigma, accessed July 2026.',
        '[3] FastAPI Documentation, "FastAPI framework, high performance, easy to learn," https://fastapi.tiangolo.com/, accessed July 2026.',
        '[4] Neo4j Inc., "Neo4j Graph Database Documentation," https://neo4j.com/docs/, accessed July 2026.',
        '[5] Apache Software Foundation, "Apache Kafka Documentation," https://kafka.apache.org/documentation/, accessed July 2026.',
        '[6] OpenSearch Project, "OpenSearch Documentation," https://opensearch.org/docs/latest/, accessed July 2026.',
        '[7] React Team, "React Documentation," https://react.dev/, accessed July 2026.',
        '[8] Docker Inc., "Docker Documentation," https://docs.docker.com/, accessed July 2026.',
        '[9] NIST, "NIST Cybersecurity Framework," https://www.nist.gov/cyberframework, accessed July 2026.',
        '[10] Redpanda Data, "Redpanda Documentation," https://docs.redpanda.com/, accessed July 2026.',
    ]
    for ref in refs:
        _add_paragraph(doc, ref, size=12, spacing=1.0, after=6)


def build_report():
    cover_path = DOWNLOADS / "Cover-Projectbased.docx"
    cert_path = DOWNLOADS / "Certificate-7th Sem.docx"

    master = Document(str(cover_path))
    composer = Composer(master)
    composer.append(Document(str(cert_path)))

    doc = master  # composer modifies master in place

    # Front matter section with margins
    _company_certificate(doc)

    # New section for roman-numbered front matter
    new_sec = doc.add_section()
    _set_margins(new_sec)
    _header_footer(new_sec, roman=True)

    _acknowledgement(doc)
    _abstract(doc)
    _add_toc(doc)
    _list_of_figures(doc)
    _list_of_tables(doc)
    _company_description(doc)

    # Chapter section (arabic numbering)
    chapter_sec = doc.add_section()
    _set_margins(chapter_sec)
    _header_footer(chapter_sec, chapter_title="AegisSOC Report")

    diagram = ASSETS / "architecture_flow.png"
    _make_architecture_diagram(diagram)

    assets_dir = Path("/Users/harshil/.cursor/projects/Users-harshil-Developer-AegisSOC/assets")
    shots = {
        "investigate": assets_dir / "Screenshot_2026-07-21_at_11.53.40_AM-65a08265-1dbc-4033-aa17-ae620ffce086.png",
        "replay": assets_dir / "Screenshot_2026-07-21_at_11.17.03_AM-2bbe3033-3371-450e-8a7c-fd4cfc2f809d.png",
    }

    _chapter1(doc)
    _chapter2(doc)
    _chapter3(doc)
    _chapter4(doc)
    _chapter5(doc, diagram)
    _chapter6(doc, shots)
    _chapter7(doc)
    _chapter8(doc)
    _chapter9(doc)
    _conclusion(doc)
    _references(doc)

    # Apply base font to Normal style
    try:
        doc.styles["Normal"].font.name = FONT
        doc.styles["Normal"].font.size = Pt(12)
    except Exception:
        pass

    doc.save(str(OUT))
    doc.save(str(OUT_DL))
    print(f"Saved: {OUT}")
    print(f"Saved: {OUT_DL}")


if __name__ == "__main__":
    build_report()
